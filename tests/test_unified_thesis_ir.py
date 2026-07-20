from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from extract_source_evidence import extract_docx_evidence, extract_text_evidence  # noqa: E402
from thesis_ir import build_thesis_ir  # noqa: E402
from validate_thesis_ir import validate  # noqa: E402


class UnifiedThesisIrTests(unittest.TestCase):
    def build_sources(self, directory: Path) -> list[Path]:
        txt = directory / "sample.txt"
        txt.write_text(
            "统一语义模型测试\n摘要\n这是一段摘要。\n关键词：论文；排版\n\n1 引言\n这是正文。\n\n参考文献\n[1] Example.\n",
            encoding="utf-8",
        )
        md = directory / "sample.md"
        md.write_text(
            "# 统一语义模型测试\n\n# 摘要\n\n这是一段摘要。\n\n**关键词：** 论文；排版\n\n# 1 引言\n\n这是正文。\n\n"
            "|列A|列B|\n|---|---|\n|1|2|\n\n![示意图](figure.png)\n\n# 参考文献\n\n[1] Example.\n",
            encoding="utf-8",
        )
        docx = directory / "sample.docx"
        document = Document()
        document.add_paragraph("统一语义模型测试")
        document.add_paragraph("摘要：这是一段摘要。")
        document.add_paragraph("关键词：论文；排版")
        document.add_heading("1 引言", level=1)
        document.add_paragraph("这是正文。")
        document.add_heading("参考文献", level=1)
        document.add_paragraph("[1] Example.")
        document.save(docx)
        return [txt, md, docx]

    def test_all_sources_converge_to_v2_schema(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            sources = self.build_sources(Path(tmp))
            formats = []
            for source in sources:
                if source.suffix == ".docx":
                    evidence = extract_docx_evidence(source, None, None)
                else:
                    evidence = extract_text_evidence(source)
                thesis_ir = build_thesis_ir(evidence)
                report = validate(thesis_ir)
                self.assertTrue(report["passed"], report["errors"])
                self.assertEqual(thesis_ir["schema"], "paper-formatter.thesis-ir")
                self.assertEqual(thesis_ir["version"], "2.0")
                self.assertGreater(thesis_ir["semanticBlockCount"], 0)
                formats.append(thesis_ir["source"]["format"])
            self.assertEqual(formats, ["text", "markdown", "docx"])

    def test_markdown_assets_are_semantic_blocks(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            md = self.build_sources(Path(tmp))[1]
            thesis_ir = build_thesis_ir(extract_text_evidence(md))
            kinds = [item["kind"] for item in thesis_ir["semanticBlocks"]]
            self.assertIn("table", kinds)
            self.assertIn("image", kinds)
            image = next(item for item in thesis_ir["semanticBlocks"] if item["kind"] == "image")
            self.assertEqual(image["attributes"]["path"], "figure.png")

    def test_docx_keeps_word_native_capabilities(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = self.build_sources(Path(tmp))[2]
            thesis_ir = build_thesis_ir(extract_docx_evidence(docx, None, None))
            capabilities = thesis_ir["source"]["capabilities"]
            self.assertTrue(capabilities["wordNativeStructure"])
            self.assertTrue(capabilities["styles"])
            self.assertTrue(capabilities["sections"])


if __name__ == "__main__":
    unittest.main()
