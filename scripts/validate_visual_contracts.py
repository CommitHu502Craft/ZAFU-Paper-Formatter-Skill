#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document


def split_label_content(paragraph_text: str) -> Tuple[Optional[str], Optional[str]]:
    if "：" in paragraph_text:
        label, content = paragraph_text.split("：", 1)
        return f"{label}：", content
    if ":" in paragraph_text:
        label, content = paragraph_text.split(":", 1)
        return f"{label}:", content
    return None, None


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def first_nonempty_run_index(paragraph, start: int = 0) -> Optional[int]:
    for index in range(start, len(paragraph.runs)):
        if paragraph.runs[index].text:
            return index
    return None


def check_label_content_bold(paragraph, expected_labels: List[str], issue_type: str) -> Optional[Dict[str, Any]]:
    text = paragraph_text(paragraph).strip()
    label, content = split_label_content(text)
    if not label or label not in expected_labels or not content.strip():
        return None
    label_len = len(label)
    consumed = 0
    label_run_indexes: List[int] = []
    for index, run in enumerate(paragraph.runs):
        run_text = run.text or ""
        if not run_text:
            continue
        if consumed < label_len:
            label_run_indexes.append(index)
        consumed += len(run_text)
        if consumed >= label_len:
            break
    if not label_run_indexes:
        return None
    content_index = first_nonempty_run_index(paragraph, label_run_indexes[-1] + 1)
    label_bold = any(paragraph.runs[index].bold for index in label_run_indexes)
    content_bold = paragraph.runs[content_index].bold if content_index is not None else None
    if label_bold and content_bold in {False, None}:
        return None
    return {
        "type": issue_type,
        "text": text,
        "labelBold": label_bold,
        "contentBold": content_bold,
    }


def explicit_page_break_count(document: Document) -> int:
    count = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for child in run._element:
                if child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
                    count += 1
    return count


def blank_paragraph_count(document: Document) -> int:
    return sum(1 for paragraph in document.paragraphs if not paragraph.text.strip())


CAPTION_RE = re.compile(
    r"^(?:(图|表)\s*(?:\d+|[一二三四五六七八九十百千]+)(?:[A-Za-z])?(?:[.．\-—](?:\d+|[一二三四五六七八九十百千]+))?|"
    r"(Figure|Fig\.?|Table)\s*[0-9A-Za-z]+(?:[.．\-—][0-9A-Za-z]+)?)\s*[:：.．]?\s*.+$",
    re.IGNORECASE,
)


def caption_style_issues(document: Document, expected_style_name: str = "ZAFU Caption") -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph_text(paragraph).strip()
        if not CAPTION_RE.match(text):
            continue
        style_name = paragraph.style.name if paragraph.style else None
        if style_name == expected_style_name:
            continue
        issues.append(
            {
                "type": "caption_style_mismatch",
                "paragraphIndex": index,
                "text": text[:160],
                "expectedStyle": expected_style_name,
                "actualStyle": style_name,
            }
        )
    return issues


def validate_visual_contracts(docx_path: str) -> Dict[str, Any]:
    document = Document(docx_path)
    issues: List[Dict[str, Any]] = []
    abstract_issue = None
    keywords_issue = None
    for paragraph in document.paragraphs[:120]:
        abstract_issue = abstract_issue or check_label_content_bold(
            paragraph,
            ["摘要：", "摘要:", "Abstract:", "ABSTRACT:"],
            "abstract_label_content_format",
        )
        keywords_issue = keywords_issue or check_label_content_bold(
            paragraph,
            ["关键词：", "关键词:", "Key words:", "Key Words:", "KEY WORDS:", "Keywords:"],
            "keywords_label_content_format",
        )
    if abstract_issue:
        issues.append(abstract_issue)
    if keywords_issue:
        issues.append(keywords_issue)

    section_count = len(document.sections)
    if section_count > 5:
        issues.append(
            {
                "type": "excessive_section_count",
                "actual": section_count,
                "expectedMax": 5,
            }
        )
    page_breaks = explicit_page_break_count(document)
    if page_breaks > 12:
        issues.append(
            {
                "type": "excessive_page_breaks",
                "actual": page_breaks,
                "expectedMax": 12,
            }
        )
    blank_count = blank_paragraph_count(document)
    if blank_count > 40:
        issues.append(
            {
                "type": "excessive_blank_paragraphs",
                "actual": blank_count,
                "expectedMax": 40,
            }
        )
    issues.extend(caption_style_issues(document))
    return {
        "sourceDocx": str(Path(docx_path).resolve()),
        "issues": issues,
        "checks": {
            "abstractLabelContentBoldContract": not any(item["type"] == "abstract_label_content_format" for item in issues),
            "keywordsLabelContentBoldContract": not any(item["type"] == "keywords_label_content_format" for item in issues),
            "captionStyleContract": not any(item["type"] == "caption_style_mismatch" for item in issues),
            "sectionCountWithinCeiling": section_count <= 5,
            "pageBreakCountWithinCeiling": page_breaks <= 12,
            "blankParagraphCountWithinCeiling": blank_count <= 40,
        },
        "metrics": {
            "sectionCount": section_count,
            "pageBreakCount": page_breaks,
            "blankParagraphCount": blank_count,
        },
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Validate user-visible visual contracts for a thesis DOCX.")
    parser.add_argument("docx", help="DOCX path")
    parser.add_argument("--output", "-o", help="Output JSON path")
    args = parser.parse_args()

    payload = validate_visual_contracts(args.docx)
    text = json.dumps(payload, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
    else:
        print(text)


if __name__ == "__main__":
    main()
