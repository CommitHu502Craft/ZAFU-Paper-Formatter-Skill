#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as docx_qn
from docx.shared import Cm, Inches, Pt
from lxml import etree
from extract_source_evidence import extract_text_evidence
from docx_ooxml import load_rules
from mathml_omml import (
    INLINE_MATH_RE,
    detect_latex_block_begin,
    detect_latex_block_end,
    extract_inline_math_latex,
    extract_standalone_formula_latex,
    latex_to_omml_xml,
    split_display_math_and_number,
)
from thesis_ir import build_thesis_ir


INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
MAIN_CHAPTER_RE = re.compile(r"^(?:\d+\s+.+|第[一二三四五六七八九十百千]+章|参考文献|致谢|附录)")
REFERENCES_HEADING_MARKDOWN_RE = re.compile(r"^#{1,3}\s*参考文献(?:[（(].*?[)）])?\s*$")
ABSTRACT_CN_HEADINGS = {"中文摘要", "摘  要", "摘要"}
ABSTRACT_EN_HEADINGS = {"English Abstract", "ABSTRACT"}
KEYWORDS_CN_RE = re.compile(r"^关键词[:：]\s*(?P<content>.+)$", re.IGNORECASE)
KEYWORDS_EN_RE = re.compile(r"^key\s*words?[:：]\s*(?P<content>.+)$", re.IGNORECASE)
TITLE_CN_RE = re.compile(r"本文以[“\"](?P<title>[^”\"]+)[”\"]为题")
TITLE_EN_RE = re.compile(r'entitled\s+[“"\'](?P<title>[^”"\']+)[”"\']', re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(r"^表\s*\d+(?:[-—.．]\d+)?")
IMAGE_MARKDOWN_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")
PLAIN_TEXT_HEADING_RE = re.compile(
    r"^(?:"
    r"\d+(?:\.\d+)*\s+.+|"
    r"[一二三四五六七八九十百千]+[、.．]\s*.+|"
    r"第[一二三四五六七八九十百千]+章.+|"
    r"第[一二三四五六七八九十百千]+节.+|"
    r"参考文献|致谢|附录.+|摘要|English Abstract|ABSTRACT"
    r")$"
)

A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
EQUATION_NUMBER_RE = re.compile(r"(?P<number>[（(]\d+(?:[-.．]\d+)*[）)])$")
NUMERIC_REFERENCE_LABEL_RE = re.compile(r"^\[\s*\d+\s*\]\s*")


def strip_markdown_emphasis(text: str) -> str:
    return (text or "").replace("**", "").replace("__", "").strip()


def normalize_english_keywords(text: Optional[str]) -> Optional[str]:
    if not text:
        return text
    value = re.sub(r"\s*[;；]\s*", ", ", text.strip())
    value = re.sub(r"\s*,\s*", ", ", value)
    return value.strip(" ,")


def normalize_inline_text(text: str) -> str:
    value = (text or "").replace("\t", "    ")
    return value.strip()


def normalize_math_source(text: str) -> str:
    value = normalize_inline_text(text)
    if not value:
        return value
    # Strip the common Markdown math fences first so the downstream
    # LaTeX-to-OMML path sees the raw expression instead of the wrapper.
    if value.startswith("$$") and value.endswith("$$") and len(value) >= 4:
        value = value[2:-2].strip()
    elif value.startswith(r"\[") and value.endswith(r"\]") and len(value) >= 4:
        value = value[2:-2].strip()
    return value


def source_looks_like_markdown(lines: List[str]) -> bool:
    return any(re.match(r"^#{1,6}\s+", line.strip()) for line in lines[:100])


def latex_to_omml(latex: str, display: str = "inline"):
    omml_xml = latex_to_omml_xml(latex, display=display)
    if omml_xml is None:
        return None
    return etree.fromstring(omml_xml)


def clear_document(document: Document) -> None:
    document._body.clear_content()


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    node = settings.find(docx_qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(docx_qn("w:val"), "true")


def set_section_page_number_format(section, start: Optional[int] = None, fmt: Optional[str] = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(docx_qn("w:pgNumType"))
    if start is None and fmt is None:
        if pg_num_type is not None:
            sect_pr.remove(pg_num_type)
        return
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    if start is not None:
        pg_num_type.set(docx_qn("w:start"), str(start))
    if fmt is not None:
        pg_num_type.set(docx_qn("w:fmt"), fmt)


def force_section_a4(section) -> None:
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)


def add_toc_entry_field(paragraph, entry_text: str, level: int, identifier: str = "A") -> None:
    run_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(docx_qn("w:fldCharType"), "begin")
    run_begin.append(fld_begin)

    run_instr = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    vanish = OxmlElement("w:vanish")
    no_proof = OxmlElement("w:noProof")
    rpr.append(vanish)
    rpr.append(no_proof)
    run_instr.append(rpr)
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f' TC "{entry_text}" \\f {identifier} \\l {level} '
    run_instr.append(instr)

    run_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(docx_qn("w:fldCharType"), "end")
    run_end.append(fld_end)

    paragraph._p.append(run_begin)
    paragraph._p.append(run_instr)
    paragraph._p.append(run_end)


def append_inline_math(paragraph, latex: str) -> bool:
    omml = latex_to_omml(latex, display="inline")
    if omml is None:
        return False
    paragraph._p.append(omml)
    return True


def add_inline_runs(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_MATH_RE.finditer(text):
        before = text[cursor : match.start()]
        if before:
            add_inline_runs_without_math(paragraph, before)
        token = match.group(0)
        latex = extract_inline_math_latex(token)
        if not append_inline_math(paragraph, latex.strip()):
            paragraph.add_run(token)
        cursor = match.end()
    tail = text[cursor:]
    if tail:
        add_inline_runs_without_math(paragraph, tail)


def add_inline_runs_without_math(paragraph, text: str) -> None:
    parts = INLINE_TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            rfonts = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
            if rfonts is not None:
                rfonts.set(docx_qn("w:ascii"), "Consolas")
                rfonts.set(docx_qn("w:hAnsi"), "Consolas")
                rfonts.set(docx_qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def ensure_page_break(document: Document, state: dict) -> None:
    if state.get("is_first_block"):
        state["pending_page_break"] = False
        return
    if state.get("last_block_was_page_break"):
        state["pending_page_break"] = False
        return
    document.add_page_break()
    state["last_block_was_page_break"] = True
    state["is_first_block"] = False
    state["pending_page_break"] = False


def consume_pending_page_break(document: Document, state: dict) -> None:
    if state.get("pending_page_break"):
        ensure_page_break(document, state)


def mark_content_written(state: dict) -> None:
    state["is_first_block"] = False
    state["last_block_was_page_break"] = False
    state["reuse_last_empty_paragraph"] = False


def acquire_output_paragraph(document: Document, state: dict, *, style: Optional[str] = None):
    reusable = bool(state.get("reuse_last_empty_paragraph"))
    if reusable and document.paragraphs:
        paragraph = document.paragraphs[-1]
        if not (paragraph.text or "").strip():
            if style:
                paragraph.style = style
            state["reuse_last_empty_paragraph"] = False
            return paragraph
    paragraph = document.add_paragraph(style=style)
    state["reuse_last_empty_paragraph"] = False
    return paragraph


def add_paragraph_block(document: Document, text: str, style: Optional[str], state: dict):
    consume_pending_page_break(document, state)
    paragraph = acquire_output_paragraph(document, state, style=style)
    add_inline_runs(paragraph, text)
    mark_content_written(state)
    return paragraph


def add_reference_paragraph_block(document: Document, text: str, state: dict):
    return add_paragraph_block(document, text, style=None, state=state)


def add_body_paragraph_block(document: Document, text: str, state: dict):
    return add_paragraph_block(document, text, style="Body Text", state=state)


def add_centered_title_block(document: Document, text: str, state: dict):
    consume_pending_page_break(document, state)
    paragraph = acquire_output_paragraph(document, state, style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(paragraph, text)
    mark_content_written(state)
    return paragraph


def add_display_math_block(document: Document, latex: str, state: dict) -> None:
    consume_pending_page_break(document, state)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    omml = latex_to_omml(normalize_math_source(latex), display="block")
    if omml is None:
        paragraph.add_run(f"$$ {latex} $$")
    else:
        paragraph._p.append(omml)
    mark_content_written(state)


def add_display_math_with_number_block(document: Document, latex: str, number: str, state: dict) -> None:
    consume_pending_page_break(document, state)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    section = document.sections[-1]
    page_width = section.page_width or Cm(A4_WIDTH_CM)
    left_margin = section.left_margin or Cm(2.7)
    right_margin = section.right_margin or Cm(2.7)
    text_width = page_width - left_margin - right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(int(text_width / 2), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(int(text_width), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t")
    omml = latex_to_omml(normalize_math_source(latex), display="inline")
    if omml is None:
        paragraph.add_run(latex)
    else:
        paragraph._p.append(omml)
    paragraph.add_run("\t")
    paragraph.add_run(number)
    mark_content_written(state)


def convert_svg_to_png(svg_path: Path) -> Optional[Path]:
    """Convert SVG to PNG using cairosvg. Returns PNG path or None if conversion fails."""
    try:
        import cairosvg
    except ImportError:
        return None

    png_path = svg_path.with_suffix(".png")
    try:
        cairosvg.svg2png(
            url=str(svg_path),
            write_to=str(png_path),
            output_width=2400,
            dpi=300,
        )
        return png_path
    except Exception:
        return None


def add_image_block(document: Document, alt_text: str, image_path: str, state: dict, markdown_dir: Path) -> None:
    consume_pending_page_break(document, state)
    resolved = (markdown_dir / image_path).resolve()
    supported = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp"}

    # Handle SVG by converting to PNG
    actual_path = resolved
    if resolved.exists() and resolved.suffix.lower() == ".svg":
        png_path = convert_svg_to_png(resolved)
        if png_path and png_path.exists():
            actual_path = png_path
        else:
            # SVG conversion failed, add placeholder
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder = alt_text or resolved.stem or Path(image_path).stem or "图形"
            paragraph.add_run(f"[图形占位] {placeholder}")
            mark_content_written(state)
            return

    if actual_path.exists() and actual_path.suffix.lower() in supported:
        document.add_picture(str(actual_path), width=Inches(5.8))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder = alt_text or resolved.stem or Path(image_path).stem or "图形"
        paragraph.add_run(f"[图形占位] {placeholder}")
    mark_content_written(state)


def add_heading_block(document: Document, text: str, level: int, state: dict):
    is_main_chapter = level == 1 and MAIN_CHAPTER_RE.match(text or "")
    if is_main_chapter and state.get("chapter_page_breaks"):
        ensure_page_break(document, state)
    else:
        consume_pending_page_break(document, state)
    paragraph = acquire_output_paragraph(document, state, style=f"Heading {level}")
    add_inline_runs(paragraph, text)
    add_toc_entry_field(paragraph, text, level=min(level, 3))
    mark_content_written(state)
    return paragraph


def add_code_block(document: Document, lines: List[str], state: dict) -> None:
    if not lines:
        return
    consume_pending_page_break(document, state)
    paragraph = document.add_paragraph()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        rfonts = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
        if rfonts is not None:
            rfonts.set(docx_qn("w:ascii"), "Consolas")
            rfonts.set(docx_qn("w:hAnsi"), "Consolas")
            rfonts.set(docx_qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(10)
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)
    mark_content_written(state)


def parse_table_row(line: str) -> List[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [cell.strip() for cell in raw.split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_table_block(document: Document, lines: List[str], state: dict) -> None:
    if not lines:
        return
    consume_pending_page_break(document, state)
    rows = [parse_table_row(line) for line in lines]
    if len(rows) >= 2 and is_table_separator(lines[1]):
        rows = [rows[0], *rows[2:]]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    def populate_table_cell(cell, text: str) -> None:
        # `cell.text = ...` strips math objects and flattens inline formatting.
        # Rebuild the cell body through the same inline parser used for正文.
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        text = normalize_inline_text(text)
        if not text:
            return
        display_latex, equation_number = split_display_math_and_number(text)
        if display_latex and display_latex == text:
            omml = latex_to_omml(display_latex, display="block")
            if omml is not None:
                paragraph._p.append(omml)
                return
        standalone_latex = extract_standalone_formula_latex(text)
        if standalone_latex and standalone_latex == text:
            if not append_inline_math(paragraph, standalone_latex):
                paragraph.add_run(text)
            return
        add_inline_runs(paragraph, text)

    for row_index, row_values in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            populate_table_cell(cell, row_values[column_index] if column_index < len(row_values) else "")
    mark_content_written(state)


def add_toc_field(paragraph) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(docx_qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "在 Word 中打开后更新目录"
    run.append(text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


def insert_toc_block(document: Document, state: dict) -> None:
    if state.get("toc_inserted"):
        return
    consume_pending_page_break(document, state)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = "Title"
    title_run = title.add_run("目  录")
    title_run.font.size = Pt(22)
    toc_paragraph = document.add_paragraph()
    add_toc_field(toc_paragraph)
    state["toc_inserted"] = True
    state["is_first_block"] = False
    state["last_block_was_page_break"] = True
    state["pending_page_break"] = False


def parse_front_abstract_block(lines: List[str]) -> Optional[dict]:
    index = 0
    while index < len(lines) and not lines[index].strip():
        index += 1
    if index >= len(lines):
        return None

    first_heading = re.match(r"^#\s+(.+)$", lines[index].strip())
    if not first_heading or normalize_inline_text(first_heading.group(1)) not in ABSTRACT_CN_HEADINGS:
        return None
    index += 1

    cn_paragraphs: List[str] = []
    keywords_cn: Optional[str] = None
    found_keywords_heading = False
    while index < len(lines):
        stripped = lines[index].strip()
        plain = strip_markdown_emphasis(stripped)
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            index += 1
            break
        # Stop if we hit the English abstract heading
        if re.match(r"^#\s+", stripped) and any(kw in stripped for kw in ["English Abstract", "ABSTRACT"]):
            break
        # Handle level 2 headings
        if re.match(r"^##\s+", stripped):
            heading_text = re.match(r"^##\s+(.+)$", stripped)
            if heading_text and "关键词" in heading_text.group(1):
                found_keywords_heading = True
                index += 1
                continue
            # Otherwise skip the heading itself
            index += 1
            continue
        # If we just found the keywords heading, this line is the keywords
        if found_keywords_heading and not keywords_cn:
            keywords_cn = normalize_inline_text(plain)
            found_keywords_heading = False
            index += 1
            continue
        match = KEYWORDS_CN_RE.match(plain)
        if match:
            keywords_cn = normalize_inline_text(match.group("content"))
            index += 1
            continue
        cn_paragraphs.append(plain)
        index += 1

    while index < len(lines) and not lines[index].strip():
        index += 1
    if index >= len(lines):
        return None

    second_heading = re.match(r"^#\s+(.+)$", lines[index].strip())
    if not second_heading or normalize_inline_text(second_heading.group(1)) not in ABSTRACT_EN_HEADINGS:
        return None
    index += 1

    en_paragraphs: List[str] = []
    keywords_en: Optional[str] = None
    found_en_keywords_heading = False
    while index < len(lines):
        stripped = lines[index].strip()
        plain = strip_markdown_emphasis(stripped)
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            index += 1
            break
        # Stop if we hit the next level 1 heading (Chapter 1, etc.)
        if re.match(r"^#\s+", stripped) and not any(kw in stripped for kw in ["English Abstract", "ABSTRACT"]):
            break
        # Handle level 2 headings
        if re.match(r"^##\s+", stripped):
            heading_text = re.match(r"^##\s+(.+)$", stripped)
            if heading_text and "keyword" in heading_text.group(1).lower():
                found_en_keywords_heading = True
                index += 1
                continue
            index += 1
            continue
        # If we just found the keywords heading, this line is the keywords
        if found_en_keywords_heading and not keywords_en:
            keywords_en = normalize_inline_text(plain)
            found_en_keywords_heading = False
            index += 1
            continue
        match = KEYWORDS_EN_RE.match(plain)
        if match:
            keywords_en = normalize_inline_text(match.group("content"))
            index += 1
            continue
        en_paragraphs.append(plain)
        index += 1

    cn_text = normalize_inline_text(" ".join(cn_paragraphs))
    en_text = normalize_inline_text(" ".join(en_paragraphs))
    if not cn_text or not en_text:
        return None

    title_cn_match = TITLE_CN_RE.search(cn_text)
    title_en_match = TITLE_EN_RE.search(en_text)
    return {
        "title_cn": normalize_inline_text(title_cn_match.group("title")) if title_cn_match else None,
        "title_en": normalize_inline_text(title_en_match.group("title")).rstrip(",.;；。") if title_en_match else None,
        "abstract_cn": cn_text,
        "abstract_en": en_text,
        "keywords_cn": keywords_cn,
        "keywords_en": normalize_english_keywords(keywords_en),
        "consumed_until": index,
    }


def emit_front_abstract_block(document: Document, block: dict, state: dict) -> None:
    if block.get("title_cn"):
        add_centered_title_block(document, block["title_cn"], state)
    abstract_cn = add_paragraph_block(document, f"摘要：{block['abstract_cn']}", style=None, state=state)
    add_toc_entry_field(abstract_cn, "摘  要", 1)
    if block.get("keywords_cn"):
        add_paragraph_block(document, f"关键词：{block['keywords_cn']}", style=None, state=state)
    if block.get("title_en"):
        add_centered_title_block(document, block["title_en"], state)
    if block.get("abstract_en"):
        abstract_en = add_paragraph_block(document, f"Abstract: {block['abstract_en']}", style=None, state=state)
        add_toc_entry_field(abstract_en, "ABSTRACT", 1)
    if block.get("keywords_en"):
        add_paragraph_block(document, f"Key words: {block['keywords_en']}", style=None, state=state)
    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    force_section_a4(body_section)
    set_section_page_number_format(body_section, start=1, fmt="decimal")
    state["last_block_was_page_break"] = True
    state["pending_page_break"] = False
    state["reuse_last_empty_paragraph"] = True


def emit_thesis_ir_text_block(document: Document, block: dict, state: dict) -> None:
    kind = block.get("kind")
    text = normalize_inline_text(str(block.get("text") or ""))
    if not text:
        return
    if kind == "heading":
        level = int(block.get("level") or 1)
        add_heading_block(document, text, max(1, min(level, 3)), state)
        return
    display_latex, equation_number = split_display_math_and_number(text)
    if display_latex:
        if equation_number:
            add_display_math_with_number_block(document, display_latex, equation_number, state)
        else:
            add_display_math_block(document, display_latex, state)
        return
    standalone_latex = extract_standalone_formula_latex(text)
    if standalone_latex and standalone_latex == text:
        add_display_math_block(document, standalone_latex, state)
        return
    if kind == "reference":
        add_reference_paragraph_block(document, text, state)
        return
    add_body_paragraph_block(document, text, state)


def normalize_reference_entry_text(text: str, thesis_ir: dict, rules: dict) -> str:
    references = thesis_ir.get("references") or {}
    normalization = references.get("normalization") or {}
    target_numbering = ((rules.get("references") or {}).get("bibliography_numbering")) or ""
    if target_numbering == "none" and normalization.get("safeListLabelStripCandidate"):
        return NUMERIC_REFERENCE_LABEL_RE.sub("", text, count=1).strip()
    return text


def normalize_reference_markdown_text(text: str, rules: dict) -> str:
    target_numbering = ((rules.get("references") or {}).get("bibliography_numbering")) or ""
    if target_numbering == "none":
        return NUMERIC_REFERENCE_LABEL_RE.sub("", text, count=1).strip()
    return text


def build_reference_entry_paragraph_text(entry: dict, thesis_ir: dict, rules: dict) -> str:
    text = normalize_reference_entry_text(str(entry.get("text") or ""), thesis_ir, rules)
    label = entry.get("label")
    if label and not NUMERIC_REFERENCE_LABEL_RE.match(text):
        text = f"[{label}] {text}"
    return text


def consume_thesis_ir_latex_block(ordered_blocks: List[dict], start_index: int) -> tuple[int, Optional[str], Optional[str]]:
    if start_index >= len(ordered_blocks):
        return start_index, None, None
    first = ordered_blocks[start_index]
    if first.get("kind") != "paragraph":
        return start_index, None, None
    env = detect_latex_block_begin(normalize_inline_text(str(first.get("text") or "")))
    if not env:
        return start_index, None, None
    latex_lines: List[str] = []
    index = start_index + 1
    while index < len(ordered_blocks):
        candidate = ordered_blocks[index]
        if candidate.get("kind") != "paragraph":
            break
        text = normalize_inline_text(str(candidate.get("text") or ""))
        suffix = detect_latex_block_end(text, env)
        if suffix is not None:
            equation_number = None
            if suffix:
                match = EQUATION_NUMBER_RE.search(suffix)
                equation_number = match.group("number") if match else None
            latex = "\n".join(line for line in latex_lines if line).strip()
            return index + 1, latex or None, equation_number
        latex_lines.append(text)
        index += 1
    return start_index, None, None


def build_plain_text_docx_from_thesis_ir(document: Document, thesis_ir: dict, rules: dict, state: dict) -> None:
    front_matter = thesis_ir.get("frontMatter") or {}
    front_block = {
        "title_cn": front_matter.get("title"),
        "title_en": front_matter.get("titleEn") or front_matter.get("title_en"),
        "abstract_cn": front_matter.get("abstractCn"),
        "abstract_en": front_matter.get("abstractEn"),
        "keywords_cn": front_matter.get("keywordsCn"),
        "keywords_en": front_matter.get("keywordsEn"),
    }
    if front_block.get("abstract_cn"):
        if not state.get("toc_inserted"):
            insert_toc_block(document, state)
        emit_front_abstract_block(document, front_block, state)

    semantic_blocks = thesis_ir.get("semanticBlocks") or []
    if semantic_blocks:
        source_path = ((thesis_ir.get("source") or {}).get("path")) or ((thesis_ir.get("evidenceSummary") or {}).get("source"))
        source_dir = Path(source_path).resolve().parent if source_path else Path.cwd()
        render_blocks: List[dict] = []
        for block in semantic_blocks:
            role = str(block.get("role") or "")
            if role == "front_matter":
                continue
            kind = str(block.get("kind") or "paragraph")
            normalized = dict(block)
            if kind == "caption":
                normalized["kind"] = "paragraph"
            elif kind == "reference":
                normalized["kind"] = "reference"
                normalized["text"] = normalize_reference_entry_text(str(block.get("text") or ""), thesis_ir, rules)
            render_blocks.append(normalized)

        index = 0
        while index < len(render_blocks):
            next_index, latex_block, equation_number = consume_thesis_ir_latex_block(render_blocks, index)
            if latex_block is not None and next_index > index:
                if equation_number:
                    add_display_math_with_number_block(document, latex_block, equation_number, state)
                else:
                    add_display_math_block(document, latex_block, state)
                index = next_index
                continue
            block = render_blocks[index]
            kind = str(block.get("kind") or "paragraph")
            attributes = block.get("attributes") or {}
            if kind == "image":
                image_path = attributes.get("path")
                if image_path:
                    add_image_block(
                        document,
                        str(attributes.get("alt") or block.get("text") or ""),
                        str(image_path),
                        state,
                        source_dir,
                    )
                index += 1
                continue
            if kind == "table":
                table_lines = attributes.get("lines") or []
                if table_lines:
                    add_table_block(document, list(table_lines), state)
                # DOCX-native tables are reattached by the hybrid asset stage.
                index += 1
                continue
            emit_thesis_ir_text_block(document, block, state)
            index += 1
        return

    ordered_blocks: List[dict] = []
    for item in thesis_ir.get("headingTree") or []:
        ordered_blocks.append(
            {
                "kind": "heading",
                "sourceIndex": item.get("sourceIndex"),
                "text": item.get("text"),
                "level": item.get("level"),
            }
        )
    for item in thesis_ir.get("captionBlocks") or []:
        ordered_blocks.append(
            {
                "kind": "paragraph",
                "sourceIndex": item.get("sourceIndex"),
                "text": item.get("text"),
            }
        )
    ordered_blocks.extend(thesis_ir.get("bodyBlocks") or [])
    acknowledgements = thesis_ir.get("acknowledgements") or {}
    if acknowledgements.get("heading") and acknowledgements.get("headingSourceIndex") is not None:
        ordered_blocks.append(
            {
                "kind": "heading",
                "sourceIndex": acknowledgements.get("headingSourceIndex"),
                "text": acknowledgements.get("heading"),
                "level": 1,
            }
        )
    for block in acknowledgements.get("blocks") or []:
        ordered_blocks.append(block)
    appendix = thesis_ir.get("appendix") or {}
    for section in appendix.get("sections") or []:
        if section.get("heading") and section.get("headingSourceIndex") is not None:
            ordered_blocks.append(
                {
                    "kind": "heading",
                    "sourceIndex": section.get("headingSourceIndex"),
                    "text": section.get("heading"),
                    "level": 1,
                }
            )
        for block in section.get("blocks") or []:
            ordered_blocks.append(block)
    references = thesis_ir.get("references") or {}
    reference_entries = references.get("entries") or []
    if references.get("heading") and references.get("headingSourceIndex") is not None:
        ordered_blocks.append(
            {
                "kind": "heading",
                "sourceIndex": references.get("headingSourceIndex"),
                "text": references.get("heading"),
                "level": 1,
            }
        )
    for entry in reference_entries:
        ordered_blocks.append(
            {
                **entry,
                "kind": "reference",
                "text": build_reference_entry_paragraph_text(entry, thesis_ir, rules),
            }
        )
    ordered_blocks.sort(key=lambda item: int(item.get("sourceIndex") or 0))

    index = 0
    while index < len(ordered_blocks):
        next_index, latex_block, equation_number = consume_thesis_ir_latex_block(ordered_blocks, index)
        if latex_block is not None and next_index > index:
            if equation_number:
                add_display_math_with_number_block(document, latex_block, equation_number, state)
            else:
                add_display_math_block(document, latex_block, state)
            index = next_index
            continue
        emit_thesis_ir_text_block(document, ordered_blocks[index], state)
        index += 1


def build_source_docx_from_markdown(markdown_path: str, output_docx: str, rules: dict, thesis_ir: Optional[dict] = None) -> None:
    lines = Path(markdown_path).read_text(encoding="utf-8").splitlines()
    markdownish = source_looks_like_markdown(lines)
    if thesis_ir and thesis_ir.get("semanticBlocks"):
        markdownish = False
    markdown_dir = Path(markdown_path).resolve().parent
    if not markdownish and thesis_ir is None:
        thesis_ir = build_thesis_ir(extract_text_evidence(Path(markdown_path)))
    document = Document()
    clear_document(document)
    set_update_fields_on_open(document)
    force_section_a4(document.sections[0])

    state = {
        "toc_inserted": False,
        "pending_page_break": False,
        "is_first_block": True,
        "last_block_was_page_break": False,
        "insert_toc_for_markdown": markdownish,
        "chapter_page_breaks": False,
        "reuse_last_empty_paragraph": False,
        "in_references_section": False,
    }
    paragraph_buffer: List[str] = []
    table_buffer: List[str] = []
    code_buffer: List[str] = []
    in_code_block = False
    in_math_block = False
    math_block_delimiter: Optional[str] = None
    math_block_env: Optional[str] = None
    math_buffer: List[str] = []
    start_index = 0
    pending_caption_after_block: Optional[str] = None
    pending_table_caption_candidate: Optional[str] = None

    if not markdownish:
        build_plain_text_docx_from_thesis_ir(document, thesis_ir, rules, state)
        document.save(output_docx)
        return

    if markdownish:
        # Skip markdown header (HTML comments, title, markdown TOC) if present.
        # These are from combined markdown sources and should not be emitted as body text.
        while start_index < len(lines):
            stripped = lines[start_index].strip()
            if not stripped:
                start_index += 1
                continue
            if stripped.startswith("<!--"):
                start_index += 1
                continue
            if stripped.startswith("> "):
                start_index += 1
                continue
            if stripped.startswith("- [") and "](#" in stripped:
                start_index += 1
                continue
            # Stop at the first real content line, even if it is a top-level title.
            # The previous logic could skip the entire document when the abstract
            # is written inline instead of under a dedicated "# 摘要" heading.
            break

        front_abstract_block = parse_front_abstract_block(lines[start_index:])
        if front_abstract_block:
            insert_toc_block(document, state)
            emit_front_abstract_block(document, front_abstract_block, state)
            start_index += int(front_abstract_block.get("consumed_until") or 0)
        elif state.get("insert_toc_for_markdown"):
            insert_toc_block(document, state)
    def flush_paragraph_buffer() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = normalize_inline_text(" ".join(item.strip() for item in paragraph_buffer if item.strip()))
        paragraph_buffer = []
        if text:
            add_paragraph_block(document, text, style=None, state=state)

    def flush_pending_table_caption_candidate() -> None:
        nonlocal pending_table_caption_candidate
        if pending_table_caption_candidate:
            add_paragraph_block(document, pending_table_caption_candidate, style=None, state=state)
            pending_table_caption_candidate = None

    def flush_table_buffer() -> None:
        nonlocal table_buffer, pending_caption_after_block
        if not table_buffer:
            return
        add_table_block(document, table_buffer, state)
        table_buffer = []
        if pending_caption_after_block:
            add_paragraph_block(document, pending_caption_after_block, style=None, state=state)
            pending_caption_after_block = None

    def flush_code_buffer() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        add_code_block(document, code_buffer, state)
        code_buffer = []

    for raw_line in lines[start_index:]:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if in_code_block:
            if stripped.startswith("```"):
                in_code_block = False
                flush_code_buffer()
            else:
                code_buffer.append(line)
            continue

        if in_math_block:
            closing_number = None
            is_closing = False
            if math_block_delimiter == "$$" and stripped.startswith("$$"):
                is_closing = stripped == "$$" or bool(EQUATION_NUMBER_RE.search(stripped[2:].strip()))
                suffix = stripped[2:].strip()
                if suffix:
                    match = EQUATION_NUMBER_RE.search(suffix)
                    closing_number = match.group("number") if match else None
            elif math_block_delimiter == r"\[" and stripped.startswith(r"\]"):
                is_closing = stripped == r"\]" or bool(EQUATION_NUMBER_RE.search(stripped[2:].strip()))
                suffix = stripped[2:].strip()
                if suffix:
                    match = EQUATION_NUMBER_RE.search(suffix)
                    closing_number = match.group("number") if match else None
            elif math_block_env:
                suffix = detect_latex_block_end(stripped, math_block_env)
                if suffix is not None:
                    is_closing = True
                    if suffix:
                        match = EQUATION_NUMBER_RE.search(suffix)
                        closing_number = match.group("number") if match else None
            if is_closing:
                in_math_block = False
                math_expr = "\n".join(part for part in math_buffer if part is not None).strip()
                math_buffer = []
                math_block_delimiter = None
                math_block_env = None
                if math_expr:
                    if closing_number:
                        add_display_math_with_number_block(document, math_expr, closing_number, state)
                    else:
                        add_display_math_block(document, math_expr, state)
            else:
                math_buffer.append(line)
            continue

        if stripped.startswith("```"):
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            in_code_block = True
            continue

        if stripped == "$$" or stripped == r"\[":
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            flush_code_buffer()
            in_math_block = True
            math_block_delimiter = stripped
            math_block_env = None
            math_buffer = []
            continue

        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) >= 4:
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            flush_code_buffer()
            add_display_math_block(document, stripped[2:-2].strip(), state)
            continue

        if stripped.startswith(r"\[") and stripped.endswith(r"\]") and len(stripped) >= 4:
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            flush_code_buffer()
            add_display_math_block(document, stripped[2:-2].strip(), state)
            continue

        latex_block_env = detect_latex_block_begin(stripped)
        if latex_block_env:
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            flush_code_buffer()
            in_math_block = True
            math_block_delimiter = None
            math_block_env = latex_block_env
            math_buffer = []
            continue

        single_line_math, equation_number = split_display_math_and_number(stripped)
        if single_line_math:
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            flush_code_buffer()
            if equation_number:
                add_display_math_with_number_block(document, single_line_math, equation_number, state)
            else:
                add_display_math_block(document, single_line_math, state)
            continue

        if table_buffer and not stripped.startswith("|"):
            flush_table_buffer()

        if stripped == "---":
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            # Markdown horizontal rules are used as visual separators in this
            # source, not as semantic page breaks. Emitting a real page break here
            # fractures正文 and produces empty-looking pages.
            continue

        image_match = IMAGE_MARKDOWN_RE.match(stripped)
        if image_match:
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            add_image_block(
                document,
                alt_text=normalize_inline_text(image_match.group("alt")),
                image_path=image_match.group("path").strip(),
                state=state,
                markdown_dir=markdown_dir,
            )
            continue

        if stripped.startswith("|"):
            if paragraph_buffer:
                candidate = normalize_inline_text(" ".join(item.strip() for item in paragraph_buffer if item.strip()))
                if TABLE_CAPTION_RE.match(candidate):
                    paragraph_buffer = []
                    pending_caption_after_block = candidate
                else:
                    flush_paragraph_buffer()
            elif pending_table_caption_candidate:
                pending_caption_after_block = pending_table_caption_candidate
                pending_table_caption_candidate = None
            table_buffer.append(stripped)
            continue

        if not stripped:
            if paragraph_buffer:
                candidate = normalize_inline_text(" ".join(item.strip() for item in paragraph_buffer if item.strip()))
                if TABLE_CAPTION_RE.match(candidate):
                    paragraph_buffer = []
                    pending_table_caption_candidate = candidate
                else:
                    flush_paragraph_buffer()
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            level = len(heading_match.group(1))
            text = normalize_inline_text(heading_match.group(2))
            add_heading_block(document, text, level, state)
            state["in_references_section"] = bool(REFERENCES_HEADING_MARKDOWN_RE.match(stripped))
            continue

        if REFERENCES_HEADING_MARKDOWN_RE.match(stripped):
            flush_pending_table_caption_candidate()
            flush_paragraph_buffer()
            flush_table_buffer()
            flush_code_buffer()
            add_heading_block(document, normalize_inline_text(re.sub(r"^#{1,3}\s*", "", stripped)), 1, state)
            state["in_references_section"] = True
            continue

        flush_pending_table_caption_candidate()
        if state.get("in_references_section"):
            add_reference_paragraph_block(document, normalize_reference_markdown_text(normalize_inline_text(line), rules), state)
        else:
            paragraph_buffer.append(line)

    flush_pending_table_caption_candidate()
    flush_paragraph_buffer()
    flush_table_buffer()
    flush_code_buffer()
    if in_math_block and math_buffer:
        add_display_math_block(document, "\n".join(math_buffer).strip(), state)

    document.save(output_docx)


def main() -> None:
    parser = argparse.ArgumentParser(description="Preflight markdown, build a thesis DOCX, then run the OOXML repair pipeline.")
    parser.add_argument("markdown", help="Input markdown file")
    parser.add_argument("--template-docx", required=True, help="Front-matter template DOCX")
    parser.add_argument("--rules-yaml", required=True, help="Rules YAML")
    parser.add_argument("--output-dir", required=True, help="Output directory")
    parser.add_argument("--keep-source-docx", action="store_true", help="Keep the intermediate source DOCX")
    parser.add_argument("--convert-svg", action="store_true", help="Convert SVG images to PNG before building")
    parser.add_argument("--thesis-ir-json", help="Use an existing unified ThesisIR instead of regenerating it")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    evidence_json = output_dir / "source_evidence.json"
    thesis_ir_json = output_dir / "thesis_ir.json"
    numbering_json = output_dir / "numbering_recovery.json"

    # Convert SVG images to PNG if requested
    if args.convert_svg:
        md_path = Path(args.markdown)
        figures_dir = md_path.parent / "figures"
        if figures_dir.exists():
            converted = 0
            for svg_file in figures_dir.glob("*.svg"):
                png_file = svg_file.with_suffix(".png")
                if not png_file.exists():
                    result = convert_svg_to_png(svg_file)
                    if result:
                        converted += 1
            if converted > 0:
                # Update markdown references from .svg to .png
                content = md_path.read_text(encoding="utf-8")
                updated_content = re.sub(r'\.svg\)', '.png)', content)
                if updated_content != content:
                    md_path.write_text(updated_content, encoding="utf-8")
                    print(f"Converted {converted} SVG files to PNG and updated markdown references")

    if args.thesis_ir_json:
        thesis_ir = json.loads(Path(args.thesis_ir_json).read_text(encoding="utf-8"))
    else:
        subprocess.run(
            [
                sys.executable,
                str(Path(__file__).with_name("extract_source_evidence.py")),
                args.markdown,
                "--output",
                str(evidence_json),
            ],
            check=True,
        )
        subprocess.run(
            [
                sys.executable,
                str(Path(__file__).with_name("recover_numbering.py")),
                str(evidence_json),
                "--output",
                str(numbering_json),
            ],
            check=True,
        )
        subprocess.run(
            [
                sys.executable,
                str(Path(__file__).with_name("thesis_ir.py")),
                "--evidence-json",
                str(evidence_json),
                "--output",
                str(thesis_ir_json),
            ],
            check=True,
        )
        thesis_ir = json.loads(thesis_ir_json.read_text(encoding="utf-8"))
    rules = load_rules(args.rules_yaml)

    preflight_command = [
        sys.executable,
        str(Path(__file__).with_name("preflight_semantic_normalization.py")),
        args.markdown,
        "--rules-yaml",
        args.rules_yaml,
        "--output",
        str(output_dir / "source_preflight_report.json"),
    ]
    if args.thesis_ir_json:
        preflight_command.extend(["--thesis-ir-json", args.thesis_ir_json])
    subprocess.run(preflight_command, check=True)

    source_docx = output_dir / "markdown_source.docx"
    build_source_docx_from_markdown(args.markdown, str(source_docx), rules, thesis_ir=thesis_ir)

    subprocess.run(
        [
            sys.executable,
            str(Path(__file__).with_name("run_minimal_demo.py")),
            str(source_docx),
            "--template-docx",
            args.template_docx,
            "--rules-yaml",
            args.rules_yaml,
            "--output-dir",
            str(output_dir),
        ],
        check=True,
    )

    if not args.keep_source_docx and source_docx.exists():
        source_docx.unlink()


if __name__ == "__main__":
    main()
